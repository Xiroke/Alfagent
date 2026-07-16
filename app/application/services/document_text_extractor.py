"""Extract plain text from uploaded company documents."""

from __future__ import annotations

import io
import logging
from pathlib import Path

from app.core.exceptions import ValidationAppError

logger = logging.getLogger(__name__)

ALLOWED_EXTENSIONS = frozenset({".txt", ".docx", ".pdf", ".xlsx"})
MAX_EXTRACTED_CHARS = 80_000


class DocumentTextExtractor:
    """Read text from txt / docx / pdf / xlsx for LLM prefill."""

    def extract(self, filename: str, content: bytes) -> str:
        ext = Path(filename).suffix.lower()
        if ext not in ALLOWED_EXTENSIONS:
            raise ValidationAppError(
                f"Неподдерживаемый формат «{ext or 'без расширения'}». "
                f"Допустимо: {', '.join(sorted(ALLOWED_EXTENSIONS))}",
            )
        if not content:
            raise ValidationAppError("Файл пустой")

        if ext == ".txt":
            text = self._extract_txt(content)
        elif ext == ".docx":
            text = self._extract_docx(content)
        elif ext == ".pdf":
            text = self._extract_pdf(content)
        else:
            text = self._extract_xlsx(content)

        cleaned = text.strip()
        if not cleaned:
            raise ValidationAppError(
                "Не удалось извлечь текст из файла. Проверьте, что документ не скан без OCR.",
            )
        if len(cleaned) > MAX_EXTRACTED_CHARS:
            logger.warning(
                "Document %s truncated from %s to %s chars",
                filename,
                len(cleaned),
                MAX_EXTRACTED_CHARS,
            )
            cleaned = cleaned[:MAX_EXTRACTED_CHARS]
        return cleaned

    @staticmethod
    def _extract_txt(content: bytes) -> str:
        for encoding in ("utf-8", "utf-8-sig", "cp1251", "latin-1"):
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        raise ValidationAppError("Не удалось определить кодировку текстового файла")

    @staticmethod
    def _extract_docx(content: bytes) -> str:
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("python-docx is not installed") from exc

        document = Document(io.BytesIO(content))
        parts: list[str] = []
        for paragraph in document.paragraphs:
            line = paragraph.text.strip()
            if line:
                parts.append(line)
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)

    @staticmethod
    def _extract_pdf(content: bytes) -> str:
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise RuntimeError("pypdf is not installed") from exc

        reader = PdfReader(io.BytesIO(content))
        pages: list[str] = []
        for page in reader.pages:
            page_text = page.extract_text() or ""
            page_text = page_text.strip()
            if page_text:
                pages.append(page_text)
        return "\n\n".join(pages)

    @staticmethod
    def _extract_xlsx(content: bytes) -> str:
        try:
            from openpyxl import load_workbook
        except ImportError as exc:
            raise RuntimeError("openpyxl is not installed") from exc

        workbook = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        blocks: list[str] = []
        try:
            for sheet in workbook.worksheets:
                sheet_lines: list[str] = []
                for row in sheet.iter_rows(values_only=True):
                    cells = [
                        str(cell).strip()
                        for cell in row
                        if cell is not None and str(cell).strip()
                    ]
                    if cells:
                        sheet_lines.append(" | ".join(cells))
                if sheet_lines:
                    blocks.append(f"### {sheet.title}\n" + "\n".join(sheet_lines))
        finally:
            workbook.close()
        return "\n\n".join(blocks)
