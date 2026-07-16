"""Generate founders' meeting protocol (.docx) from registration draft.

Structure follows the sample in the repo:
`obrazec-protokola-obshhego-sobranija-uchreditelej.doc`
(new DOCX is created on each request — the .doc sample is a reference only).
"""

from __future__ import annotations

import io
import re
from datetime import date
from decimal import Decimal, ROUND_HALF_UP

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.schemas.wizard_prefill import (
    PrefillAddressOut,
    PrefillCompanyOut,
    PrefillFounderOut,
    PrefillTaxOut,
)

_ONES = (
    "",
    "один",
    "два",
    "три",
    "четыре",
    "пять",
    "шесть",
    "семь",
    "восемь",
    "девять",
)
_ONES_F = (
    "",
    "одна",
    "две",
    "три",
    "четыре",
    "пять",
    "шесть",
    "семь",
    "восемь",
    "девять",
)
_TEENS = (
    "десять",
    "одиннадцать",
    "двенадцать",
    "тринадцать",
    "четырнадцать",
    "пятнадцать",
    "шестнадцать",
    "семнадцать",
    "восемнадцать",
    "девятнадцать",
)
_TENS = (
    "",
    "",
    "двадцать",
    "тридцать",
    "сорок",
    "пятьдесят",
    "шестьдесят",
    "семьдесят",
    "восемьдесят",
    "девяносто",
)
_HUNDREDS = (
    "",
    "сто",
    "двести",
    "триста",
    "четыреста",
    "пятьсот",
    "шестьсот",
    "семьсот",
    "восемьсот",
    "девятьсот",
)


def _triple_to_words(n: int, feminine: bool = False) -> str:
    ones = _ONES_F if feminine else _ONES
    h, rem = divmod(n, 100)
    parts: list[str] = []
    if h:
        parts.append(_HUNDREDS[h])
    if 10 <= rem <= 19:
        parts.append(_TEENS[rem - 10])
    else:
        t, o = divmod(rem, 10)
        if t:
            parts.append(_TENS[t])
        if o:
            parts.append(ones[o])
    return " ".join(parts)


def rubles_in_words(amount: float | Decimal) -> str:
    """Format amount like: 10 000 (десять тысяч) рублей 00 копеек."""
    value = Decimal(str(amount)).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    rubles = int(value)
    kopecks = int((value - rubles) * 100)

    if rubles == 0:
        words = "ноль"
    else:
        chunks: list[str] = []
        millions, rem = divmod(rubles, 1_000_000)
        thousands, units = divmod(rem, 1_000)
        if millions:
            w = _triple_to_words(millions)
            form = _plural(millions, "миллион", "миллиона", "миллионов")
            chunks.append(f"{w} {form}".strip())
        if thousands:
            w = _triple_to_words(thousands, feminine=True)
            form = _plural(thousands, "тысяча", "тысячи", "тысяч")
            chunks.append(f"{w} {form}".strip())
        if units or not chunks:
            chunks.append(_triple_to_words(units) or "ноль")
        words = " ".join(c for c in chunks if c)

    rub_form = _plural(rubles, "рубль", "рубля", "рублей")
    formatted = f"{rubles:,}".replace(",", " ")
    return f"{formatted} ({words}) {rub_form} {kopecks:02d} копеек"


def _plural(n: int, one: str, few: str, many: str) -> str:
    n = abs(n) % 100
    if 11 <= n <= 14:
        return many
    last = n % 10
    if last == 1:
        return one
    if 2 <= last <= 4:
        return few
    return many


def _short_name(full: str) -> str:
    parts = [p for p in re.split(r"\s+", full.strip()) if p]
    if not parts:
        return "—"
    if len(parts) == 1:
        return parts[0]
    initials = "".join(f"{p[0]}." for p in parts[1:] if p)
    return f"{parts[0]} {initials}".strip()


def _company_full(name: str) -> str:
    cleaned = name.strip()
    if not cleaned:
        return "Общество с ограниченной ответственностью «________»"
    low = cleaned.lower()
    if "общество с ограниченной ответственностью" in low or cleaned.upper().startswith("ООО"):
        return cleaned
    return f"Общество с ограниченной ответственностью «{cleaned}»"


def _company_short(name: str, short: str) -> str:
    if short.strip():
        s = short.strip()
        return s if s.upper().startswith("ООО") else f"ООО «{s}»"
    cleaned = name.strip().strip("«»\"'")
    if cleaned.upper().startswith("ООО"):
        return cleaned
    return f"ООО «{cleaned or '________'}»"


def _compose_address(address: PrefillAddressOut) -> str:
    if address.full_address.strip():
        return address.full_address.strip()
    parts = [
        address.postal_code,
        address.region,
        address.city,
        address.street,
        f"д. {address.building}" if address.building else "",
        f"кв. {address.apartment}" if address.apartment else "",
    ]
    return ", ".join(p for p in parts if p) or "________________"


def _founder_line(f: PrefillFounderOut, address: PrefillAddressOut) -> str:
    bits = [f.full_name or "________________"]
    if f.inn:
        bits.append(f"ИНН {f.inn}")
    if f.phone:
        bits.append(f"тел. {f.phone}")
    if f.email:
        bits.append(f"email {f.email}")
    # Use company legal address as residence fallback when home registration
    if address.address_type.value == "home" and address.full_address:
        bits.append(f"место жительства: {address.full_address}")
    return "; ".join(bits)


class ProtocolDocumentService:
    def build_docx_bytes(
        self,
        *,
        company: PrefillCompanyOut,
        founders: list[PrefillFounderOut],
        address: PrefillAddressOut,
        tax: PrefillTaxOut,
        meeting_date: date | None = None,
    ) -> bytes:
        if not founders:
            founders = [
                PrefillFounderOut(
                    full_name="________________",
                    ownership_share=100,
                    is_director=True,
                )
            ]

        meeting_date = meeting_date or date.today()
        date_str = meeting_date.strftime("%d.%m.%Y")
        city = address.city.strip() or "Москва"
        full_name = _company_full(company.name)
        short_name = _company_short(company.name, company.short_name)
        capital = Decimal(str(company.authorized_capital or 10_000))
        capital_words = rubles_in_words(capital)
        location = _compose_address(address)

        chair = next((f for f in founders if f.is_director), founders[0])
        secretary = founders[1] if len(founders) > 1 else founders[0]
        vote_counter = secretary

        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)

        def center(text: str, *, bold: bool = False) -> None:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.bold = bold
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

        def body(text: str, *, bold: bool = False) -> None:
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = bold
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

        center("ПРОТОКОЛ", bold=True)
        center("Общего собрания учредителей", bold=True)
        center(full_name, bold=True)
        center("(далее — Общество)")
        center(f"город {city}")
        center(date_str)

        body(f"Дата проведения общего собрания: {date_str}")
        body("Форма проведения общего собрания: собрание (совместное присутствие)")
        body(f"Место проведения общего собрания: {location}")
        body(f"Дата составления протокола общего собрания: {date_str}")
        body(f"Подсчёт голосов произвёл: {vote_counter.full_name or '________________'}")
        body(f"Председательствующий на общем собрании: {chair.full_name or '________________'}")
        body(f"Секретарь общего собрания: {secretary.full_name or '________________'}")
        body(f"Всего учредителей Общества, включённых в списки для голосования: {len(founders)}")

        body("Учредители:", bold=True)
        for f in founders:
            body(f"— {_founder_line(f, address)}")
        body(f"Итого: {len(founders)} {_plural(len(founders), 'учредитель', 'учредителя', 'учредителей')}")

        body(
            "На общем собрании присутствуют все учредители Общества, кворум (100%) в наличии, "
            "общее собрание правомочно принимать решения по всем вопросам повестки дня."
        )

        body("Повестка дня", bold=True)
        agenda = [
            "1. Избрание председательствующего и секретаря общего собрания учредителей и возложение обязанности по подсчёту голосов.",
            "2. Учреждение Общества с ограниченной ответственностью.",
            "3. Утверждение фирменного наименования Общества.",
            "4. Утверждение размера уставного капитала Общества, а также порядка, способа и сроков образования имущества Общества.",
            "5. Утверждение размера и номинальной стоимости долей учредителей Общества.",
            "6. Утверждение места нахождения Общества.",
            "7. Заключение договора об учреждении Общества.",
            "8. Утверждение Устава Общества.",
            "9. Избрание Генерального директора Общества.",
            "10. Определение порядка совместной деятельности учредителей по созданию Общества и осуществлению государственной регистрации Общества.",
            "11. Оплата государственной пошлины за государственную регистрацию Общества.",
            "12. Утверждение эскиза печати Общества с назначением ответственного за изготовление и хранение печати.",
        ]
        for item in agenda:
            body(item)

        body("РЕШИЛИ", bold=True)

        body("По первому вопросу повестки дня")
        body(
            f"Избрать председательствующим на общем собрании учредителей Общества "
            f"{chair.full_name or '________________'} (далее — Председательствующий), "
            f"секретарём общего собрания учредителей Общества "
            f"{secretary.full_name or '________________'} (далее — Секретарь). "
            f"Возложить обязанность по подсчёту голосов на Секретаря общего собрания "
            f"{vote_counter.full_name or '________________'}."
        )
        body("Результаты голосования: единогласно. Решение принято единогласно.")

        body("По второму вопросу повестки дня")
        body("Учредить Общество с ограниченной ответственностью.")
        body("Результаты голосования: единогласно. Решение принято единогласно.")

        body("По третьему вопросу повестки дня")
        body(f"Утвердить полное фирменное наименование Общества на русском языке: {full_name}.")
        body(f"Утвердить сокращённое фирменное наименование Общества на русском языке: {short_name}.")
        if company.okved_codes:
            body(f"Основные коды ОКВЭД: {', '.join(company.okved_codes)}.")
        body("Результаты голосования: единогласно. Решение принято единогласно.")

        body("По четвёртому вопросу повестки дня")
        body(
            f"Утвердить уставный капитал Общества в размере {capital_words}, что составляет 100%. "
            f"Оплата производится денежными средствами в размере {capital_words}. "
            f"На момент государственной регистрации Общества уставный капитал оплачивается в размере 0,00 рублей. "
            f"100% суммы уставного капитала в размере {capital_words} будет оплачено в течение 4 (четырёх) месяцев "
            f"с даты государственной регистрации Общества."
        )
        body("Результаты голосования: единогласно. Решение принято единогласно.")

        body("По пятому вопросу повестки дня")
        body("Утвердить размер и номинальную стоимость долей учредителей Общества в следующем порядке:")
        for f in founders:
            share = Decimal(str(f.ownership_share or 0))
            nominal = (capital * share / Decimal("100")).quantize(
                Decimal("0.01"),
                rounding=ROUND_HALF_UP,
            )
            body(
                f"Учредитель {f.full_name or '________________'}: "
                f"номинальная стоимость доли {rubles_in_words(nominal)}; "
                f"размер доли {share}%."
            )
        body("Результаты голосования: единогласно. Решение принято единогласно.")

        body("По шестому вопросу повестки дня")
        body(
            f"Утвердить место нахождения Общества (место нахождения его постоянно действующего "
            f"исполнительного органа): {location}."
        )
        body("Результаты голосования: единогласно. Решение принято единогласно.")

        body("По седьмому вопросу повестки дня")
        body("Заключить договор об учреждении Общества.")
        body("Результаты голосования: единогласно. Решение принято единогласно.")

        body("По восьмому вопросу повестки дня")
        body("Утвердить Устав Общества.")
        body("Результаты голосования: единогласно. Решение принято единогласно.")

        body("По девятому вопросу повестки дня")
        body(
            f"Избрать Генеральным директором Общества {_founder_line(chair, address)} бессрочно. "
            f"Поручить Председательствующему подписать от имени Общества трудовой договор "
            f"с Генеральным директором после осуществления государственной регистрации."
        )
        tax_label = {"osn": "ОСН", "usn": "УСН", "ausn": "АУСН"}.get(
            tax.tax_regime.value,
            tax.tax_regime.value,
        )
        body(f"Планируемый налоговый режим: {tax_label}.")
        body("Результаты голосования: единогласно. Решение принято единогласно.")

        body("По десятому вопросу повестки дня")
        body(
            "Зарегистрировать Общество и Устав Общества в установленном законом порядке. "
            "Все действия, связанные с регистрацией Общества, а также действия, необходимые "
            "для начала деятельности Общества, возлагаются на Председательствующего. "
            "Если Общество не будет зарегистрировано, расходы компенсируются пропорционально "
            "долям учредителей в уставном капитале Общества."
        )
        body("Результаты голосования: единогласно. Решение принято единогласно.")

        body("По одиннадцатому вопросу повестки дня")
        body(
            "Поручить Председательствующему оплатить государственную пошлину за государственную "
            "регистрацию юридического лица от своего имени за всех учредителей."
        )
        body("Результаты голосования: единогласно. Решение принято единогласно.")

        body("По двенадцатому вопросу повестки дня")
        body(
            f"Утвердить эскиз печати Общества. Назначить ответственным за изготовление печати "
            f"Генерального директора Общества {chair.full_name or '________________'}."
        )
        body("Результаты голосования: единогласно. Решение принято единогласно.")

        doc.add_paragraph()
        body(f"Председательствующий ______________ / {_short_name(chair.full_name or '________')} /")
        doc.add_paragraph()
        body(f"Секретарь ______________ / {_short_name(secretary.full_name or '________')} /")

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
